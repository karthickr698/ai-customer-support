from io import BytesIO

from docx import Document

from app.application.ports.document_parser_port import DocumentParserPort, ParseSource
from app.domain.errors import DocumentParseError, EmptyDocumentError
from app.domain.ingestion import ParsedDocument, checksum_bytes
from app.rag.chunking import normalize_document_text


class DocxParserAdapter(DocumentParserPort):
    async def parse(self, source: ParseSource) -> ParsedDocument:
        if not source.binary:
            raise DocumentParseError("DOCX bytes are required")
        try:
            document = Document(BytesIO(source.binary))
            parts: list[str] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(text)
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
        except Exception as exc:  # noqa: BLE001 — python-docx raises various parse errors
            raise DocumentParseError("The DOCX document could not be parsed") from exc

        text = normalize_document_text("\n\n".join(parts))
        if not text:
            raise EmptyDocumentError("The DOCX document did not contain extractable text")

        return ParsedDocument(
            kind="docx",
            title=source.title,
            text=text,
            parser="python-docx",
            source_uri=source.source_uri,
            media_type=source.media_type
            or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            checksum=checksum_bytes(source.binary),
            character_count=len(text),
        )
